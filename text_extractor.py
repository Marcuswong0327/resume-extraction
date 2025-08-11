import os
import fitz  # PyMuPDF
import docx2txt
from docx import Document
import logging

class TextExtractor:
    """Extract text from various document formats"""
    
    def extract_text(self, file_path):
        """Extract text from PDF, DOC, or DOCX files"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.doc':
                return self._extract_from_doc(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF using PyMuPDF"""
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text.strip()
        except Exception as e:
            logging.error(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX using python-docx"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text).strip()
        except Exception as e:
            logging.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_from_doc(self, file_path):
        """Extract text from DOC using docx2txt"""
        try:
            text = docx2txt.process(file_path)
            return text.strip() if text else ""
        except Exception as e:
            logging.error(f"Error extracting DOC text: {e}")
            return ""
