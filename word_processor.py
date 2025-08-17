import docx
import streamlit as st
from io import BytesIO
import logging

class WordProcessor:
    """Enhanced Word document (.docx) text extraction with validation"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text_from_docx(self, file_content):
        """
        Extract text from DOCX file with enhanced error handling.
        
        Args:
            file_content: File content as bytes
            
        Returns:
            Extracted text as string
        """
        try:
            # Load the document from bytes
            doc = docx.Document(BytesIO(file_content))
            text_content = []

            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1

            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            text_content.append(cell.text.strip())
                            table_count += 1

            # Extract text from headers and footers
            header_footer_count = 0
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                            header_footer_count += 1
                
                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                            header_footer_count += 1

            extracted_text = "\n".join(text_content)
            
            self.logger.info(f"DOCX extraction stats - Paragraphs: {paragraph_count}, "
                           f"Table cells: {table_count}, Headers/Footers: {header_footer_count}")
            
            return extracted_text

        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise e

    def process_word_file(self, uploaded_file):
        """
        Process DOCX file and extract text with comprehensive validation.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text as string
        """
        try:
            # Validate file
            if uploaded_file.size == 0:
                self.logger.error(f"Word file {uploaded_file.name} is empty")
                st.error(f"❌ File {uploaded_file.name} is empty")
                return ""
                
            if uploaded_file.size > 10 * 1024 * 1024:  # 10MB limit for Word docs
                self.logger.error(f"Word file {uploaded_file.name} is too large ({uploaded_file.size} bytes)")
                st.error(f"❌ File {uploaded_file.name} is too large (max 10MB)")
                return ""

            uploaded_file.seek(0) 
            file_extension = uploaded_file.name.lower().split('.')[-1]

            if file_extension not in ["docx", "doc"]:
                self.logger.error(f"Unsupported file format: {file_extension}")
                st.error(f"❌ Unsupported file format: {file_extension}")
                return ""

            # Handle .doc files (older format)
            if file_extension == "doc":
                st.warning(f"⚠️ {uploaded_file.name} is in older .doc format. Please convert to .docx for better extraction.")
                return ""

            file_content = uploaded_file.read()
            
            if len(file_content) == 0:
                self.logger.error(f"No content read from {uploaded_file.name}")
                st.error(f"❌ No content could be read from {uploaded_file.name}")
                return ""

            extracted_text = self.extract_text_from_docx(file_content)
            
            if not extracted_text or not extracted_text.strip():
                self.logger.warning(f"No text extracted from {uploaded_file.name}")
                st.warning(f"⚠️ No readable text found in {uploaded_file.name}")
                return ""
                
            self.logger.info(f"Successfully extracted {len(extracted_text)} characters from {uploaded_file.name}")
            return extracted_text

        except Exception as e:
            self.logger.error(f"Error processing Word file {uploaded_file.name}: {str(e)}")
            st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
            return ""
